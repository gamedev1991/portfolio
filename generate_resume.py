from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def spacing(para, before=0, after=2, line=None):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    if line:
        para.paragraph_format.line_spacing = Pt(line)

def bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_header(title):
    p = doc.add_paragraph()
    spacing(p, before=10, after=3)
    r = p.add_run(title.upper())
    font(r, size=11, bold=True)
    bottom_border(p)

def bullet(text):
    p = doc.add_paragraph()
    spacing(p, before=1, after=2)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run('- ' + text)
    font(r, size=10)

def job_header(title, company, dates):
    p = doc.add_paragraph()
    spacing(p, before=7, after=2)
    r = p.add_run(title)
    font(r, size=10, bold=True)
    r = p.add_run('  |  ' + company)
    font(r, size=10)
    r = p.add_run('  |  ' + dates)
    font(r, size=10, italic=True)

# ── HEADER ────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p, before=0, after=2)
r = p.add_run('Rahul Ohri')
font(r, size=20, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p, before=0, after=1)
r = p.add_run('Product Manager')
font(r, size=11, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p, before=0, after=1)
r = p.add_run('+91 9742028331  |  rahulohri2007@gmail.com  |  Pune, India')
font(r, size=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p, before=0, after=6)
r = p.add_run('Portfolio: www.rahulohri.com  |  LinkedIn: [Your LinkedIn URL]')
font(r, size=10)

# ── SUMMARY ───────────────────────────────────────────────────────────────────

section_header('Summary')
for b in [
    'Launched a F2P mobile game from zero to 50,000 downloads in 5 months and nearly doubled D1 retention (10% to 18%) within two months through rapid core loop iteration.',
    'Caught a $20,000 revenue exposure mid-live-event through proactive impact analysis of prize configurations - neutralized before it reached players.',
    'Shipped 3 games in 6 months as sole PM.',
    'Vibe-coded a CSV batch automation that cut manual ops time 45x (15 min to 20 sec) - adopted team-wide.',
    'Built 11 years of gaming experience across 6.5 in engineering and 4.5 in product management - read code, owned root cause analysis, and bridged product and engineering without hand-offs.',
]:
    bullet(b)

# ── WORK EXPERIENCE ───────────────────────────────────────────────────────────

section_header('Work Experience')

job_header('Product Manager', 'JetSynthesys (Game Development Services Company)', '09/2025 - Present')
for b in [
    'Owned end-to-end LiveOps for two consecutive seasons - managing 10 production deployments, 6 LTBs, 4 Trading Posts, and full Legendary, Mythic, Rider, and Atlas branch suites on iOS and Android.',
    'Caught a $20,000 revenue exposure mid-live-event through proactive impact analysis of Tower Instance configurations during a Fortification event prize programme - neutralized before it reached players.',
    'Resolved a monetization-critical gacha defect (missing Festive Dragon shards from Draconic Chests) affecting IAP chest revenue - root cause identified and deployed to production without escalation.',
    'Vibe-coded a CSV batch automation delivering a 45x reduction in manual processing time (15 min to 20 sec) - adopted team-wide.',
    'Conducted structured impact analysis on a proposed XP Booster feature - identified unmitigable player progression risks and recommended against implementation, preventing a disruptive release from entering the development pipeline.',
]:
    bullet(b)

job_header('Product Manager', 'Blkbox AI (Creative Intelligence for Mobile Games)', '02/2024 - 04/2025')
for b in [
    'Built and led a cross-functional team from scratch - recruited a developer, artist, and game designer to develop a Free-to-Play gate runner mobile game.',
    'Shipped from greenlight to soft launch in 5 months across iOS and Android; reached 50,000 installs with a $0.80 CPI.',
    'Improved player D1 retention from 10% to 18% within two months by refining game controls and enhancing level design.',
    'Analyzed and optimized core loop and level design, increasing level completion rates by 7%.',
]:
    bullet(b)

job_header('Product Manager', 'AURA (Web3, Gaming NFT Marketplace)', '05/2023 - 11/2023')
for b in [
    'Shipped Creator Profile feature enabling user-generated NFT collections; improved featured games CTR by 5% and user engagement time by 7% through UX optimization.',
    'Implemented sprint system and streamlined Jira workflow, improving tech team delivery cadence.',
]:
    bullet(b)

job_header('Product Manager', 'FanClash (RMG, Esports Fantasy)', '04/2022 - 04/2023')
for b in [
    'Led cross-functional pod of 10+ members across design, engineering, and marketing.',
    'Revamped team creation flow, increasing conversion rate by 5%.',
    'Developed and launched Fantasy With EStars in partnership with esports teams, achieving a 90% fill rate.',
    'Improved booster engagement by 10% through user research and UX iteration.',
    'Increased adoption of Expert Opinion and Player Stats features by 15% through funnel analysis and entry point repositioning.',
]:
    bullet(b)

job_header('Game Product Manager', 'Playshifu (EdTech, STEM-based AR Toys)', '06/2021 - 03/2022')
for b in [
    'Owned full product/game lifecycle as sole PM - shipped 3 games in 6 months across design, art, and engineering.',
    'Diagnosed drop-offs in Tacto Electronics Game using Mixpanel funnel analysis; increased level completion rate by 5%.',
    'Built Mixpanel dashboard with funnel tracking for drop-off analysis and iteration strategy.',
    'Created PRDs for Plugo Animals game features, ensuring cross-functional alignment.',
]:
    bullet(b)

job_header('Senior Software Engineer', 'SplashLearn (EdTech, App-based Learning)', '07/2019 - 06/2021')
for b in [
    'Designed, developed, and launched English games on Android, iOS, and Web.',
    'Drove a 60% increase in engagement with the English Curriculum across multiple platforms.',
    'Implemented MathFacts feature, enhancing user experience and learning outcomes.',
]:
    bullet(b)

job_header('Software Engineer', 'Suventure Services Pvt Ltd. (Service-based, Finland-based Client)', '07/2018 - 07/2019')
for b in [
    'Built features for an indoor cycle training application.',
    'Contributed to Bluetooth connectivity feature, improving user experience.',
    'Developed Nearby Players List, increasing user engagement by 5%.',
]:
    bullet(b)

job_header('Game Developer', 'Nukebox Studios (F2P, Mobile Games Studio)', '12/2014 - 06/2018')
for b in [
    "Designed and developed architecture for Editor's Choice Game: Food Truck Chef.",
    'Led development of Food Truck Chef on Facebook GameRoom.',
    'Drove game improvements through player feedback analysis and competitive analysis.',
]:
    bullet(b)

# ── EDUCATION ─────────────────────────────────────────────────────────────────

section_header('Education')
for b in [
    'Product Management Certification - Upgrad (Duke CE Partnership) | 05/2020 - 11/2020',
    'Diploma in Game Programming - Asian Institute of Gaming and Animation | 05/2014 - 12/2014',
    'B.Tech in Electronics and Communication Engineering - Amity School of Engineering and Technology | 05/2010 - 04/2014',
]:
    bullet(b)

# ── SKILLS ────────────────────────────────────────────────────────────────────

section_header('Skills')
p = doc.add_paragraph()
spacing(p, before=2, after=2)
r = p.add_run(
    'LiveOps | F2P Monetization | A/B Testing | Mobile Game Development | Agile / Scrum | '
    'Product Strategy | Roadmap Planning | PRD / Product Requirements Document | Game Economy | '
    'Player Retention | KPI Tracking | DAU / MAU / Retention Metrics | Stakeholder Management | '
    'User Research | Competitive Analysis | Mixpanel | CPI / UA Metrics | OKRs | '
    'Go-to-Market Strategy | Data Analysis | Jira | Product Lifecycle Management | '
    'Sprint Planning | IAP | Unity'
)
font(r, size=10)

# ── SAVE ──────────────────────────────────────────────────────────────────────

output = 'C:/Users/RAHUL/Desktop/Portfolio/public/RahulOhri.docx'
doc.save(output)
print(f'Saved: {output}')
